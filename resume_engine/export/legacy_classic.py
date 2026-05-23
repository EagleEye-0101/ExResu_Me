"""Legacy centered single-column PDF/DOCX (classic template implementation)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListItem, ListFlowable, Paragraph, SimpleDocTemplate, Spacer

from resume_engine.schemas.resume import ResumeData
from resume_engine.utils.phone import format_phone_display


def export_pdf_classic(resume: ResumeData, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle", parent=styles["Heading1"], fontSize=16, alignment=TA_CENTER, spaceAfter=4
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontSize=9, alignment=TA_CENTER, spaceAfter=8
    )
    headline_style = ParagraphStyle(
        "Headline",
        parent=styles["Normal"],
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName="Helvetica-Bold",
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10, leading=13, alignment=TA_LEFT
    )
    job_title_style = ParagraphStyle(
        "JobTitle", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", spaceBefore=4
    )

    story = [Paragraph(resume.full_name, title_style)]
    contact_parts = [
        resume.email,
        format_phone_display(resume.phone_country_code, resume.phone),
    ]
    if resume.location:
        contact_parts.append(resume.location)
    if resume.linkedin:
        contact_parts.append(resume.linkedin)
    story.append(Paragraph(" | ".join(contact_parts), contact_style))
    if resume.headline:
        story.append(Paragraph(resume.headline, headline_style))
    if resume.summary:
        story.append(Paragraph("SUMMARY", section_style))
        story.append(Paragraph(resume.summary.replace("\n", "<br/>"), body_style))
    if resume.experience:
        story.append(Paragraph("EXPERIENCE", section_style))
        for exp in resume.experience:
            story.append(Paragraph(f"{exp.title} — {exp.company}", job_title_style))
            date_loc = f"{exp.start_date} – {exp.end_date}"
            if exp.location:
                date_loc += f" | {exp.location}"
            story.append(Paragraph(date_loc, body_style))
            if exp.bullets:
                items = [ListItem(Paragraph(b, body_style)) for b in exp.bullets]
                story.append(ListFlowable(items, bulletType="bullet", leftIndent=12))
    if resume.education:
        story.append(Paragraph("EDUCATION", section_style))
        for edu in resume.education:
            line = f"{edu.degree}"
            if edu.field:
                line += f" in {edu.field}"
            line += f" — {edu.institution}"
            if edu.graduation_date:
                line += f" ({edu.graduation_date})"
            story.append(Paragraph(line, body_style))
    if resume.skills:
        story.append(Paragraph("SKILLS", section_style))
        story.append(Paragraph(", ".join(resume.skills), body_style))
    if resume.certifications:
        story.append(Paragraph("CERTIFICATIONS", section_style))
        for cert in resume.certifications:
            line = cert.name
            if cert.issuer:
                line += f" — {cert.issuer}"
            if cert.date:
                line += f" ({cert.date})"
            story.append(Paragraph(line, body_style))
    story.append(Spacer(1, 0.2 * inch))
    doc.build(story)
    return output_path


def export_docx_classic(resume: ResumeData, output_path: Path) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume.full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_parts = [
        resume.email,
        format_phone_display(resume.phone_country_code, resume.phone),
    ]
    if resume.location:
        contact_parts.append(resume.location)
    if resume.linkedin:
        contact_parts.append(resume.linkedin)
    contact = doc.add_paragraph(" | ".join(contact_parts))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in contact.runs:
        run.font.size = Pt(10)

    if resume.headline:
        headline = doc.add_paragraph(resume.headline)
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in headline.runs:
            run.bold = True
            run.font.size = Pt(11)

    def add_section_header(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    if resume.summary:
        add_section_header("Summary")
        summary = doc.add_paragraph(resume.summary)
        for run in summary.runs:
            run.font.size = Pt(10)

    if resume.experience:
        add_section_header("Experience")
        for exp in resume.experience:
            title_line = doc.add_paragraph()
            t_run = title_line.add_run(f"{exp.title} — {exp.company}")
            t_run.bold = True
            t_run.font.size = Pt(10)
            date_line = doc.add_paragraph(f"{exp.start_date} – {exp.end_date}")
            if exp.location:
                date_line.add_run(f" | {exp.location}")
            for run in date_line.runs:
                run.font.size = Pt(9)
            for bullet in exp.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(10)

    if resume.education:
        add_section_header("Education")
        for edu in resume.education:
            line = f"{edu.degree}"
            if edu.field:
                line += f" in {edu.field}"
            line += f" — {edu.institution}"
            if edu.graduation_date:
                line += f" ({edu.graduation_date})"
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)

    if resume.skills:
        add_section_header("Skills")
        skills_p = doc.add_paragraph(", ".join(resume.skills))
        for run in skills_p.runs:
            run.font.size = Pt(10)

    if resume.certifications:
        add_section_header("Certifications")
        for cert in resume.certifications:
            line = cert.name
            if cert.issuer:
                line += f" — {cert.issuer}"
            if cert.date:
                line += f" ({cert.date})"
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
