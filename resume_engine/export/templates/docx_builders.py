"""Per-template DOCX builders — each produces visibly distinct Word output."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from resume_engine.export.templates.base import contact_pairs
from resume_engine.export.templates.ats_rules import (
    FONT_BODY_PT,
    FONT_NAME_COMPACT_PT,
    FONT_NAME_PT,
    FONT_SECTION_PT,
    FONT_SMALL_PT,
    MODERN_ACCENT_RGB,
    MARGIN_COMPACT_INCH,
    MARGIN_INCH,
)
from resume_engine.schemas.resume import ResumeData
from resume_engine.utils.phone import format_phone_display


def _margins(doc: Document, inches: float = MARGIN_INCH) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _section_header(doc: Document, title: str, color: RGBColor | None = None, underline: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(FONT_SECTION_PT)
    if color:
        run.font.color.rgb = color
    if underline:
        p.paragraph_format.border_bottom = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_bullets(doc: Document, bullets: list[str], size: int = FONT_BODY_PT, max_count: int | None = None) -> None:
    items = bullets[:max_count] if max_count else bullets
    for bullet in items:
        bp = doc.add_paragraph(bullet, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(size)


def _contact_inline(resume: ResumeData) -> str:
    parts = [
        resume.email,
        format_phone_display(resume.phone_country_code, resume.phone),
    ]
    if resume.linkedin:
        parts.append(resume.linkedin)
    if resume.github:
        parts.append(resume.github)
    if resume.location:
        parts.append(resume.location)
    return " | ".join(parts)


def build_classic_docx(resume: ResumeData, output_path: Path) -> Path:
    from resume_engine.export.legacy_classic import export_docx_classic

    return export_docx_classic(resume, output_path)


def build_professional_docx(resume: ResumeData, output_path: Path) -> Path:
    doc = Document()
    _margins(doc)

    name = doc.add_paragraph()
    nr = name.add_run(resume.full_name)
    nr.bold = True
    nr.font.size = Pt(FONT_NAME_PT)

    pairs = contact_pairs(resume)
    if pairs:
        table = doc.add_table(rows=(len(pairs) + 1) // 2, cols=2)
        table.autofit = True
        for idx, (label, value) in enumerate(pairs):
            row, col = divmod(idx, 2)
            cell = table.rows[row].cells[col]
            cell.text = ""
            p = cell.paragraphs[0]
            lr = p.add_run(f"{label}: ")
            lr.bold = True
            lr.font.size = Pt(FONT_SMALL_PT)
            vr = p.add_run(value)
            vr.font.size = Pt(FONT_SMALL_PT)

    if resume.summary:
        sp = doc.add_paragraph(resume.summary)
        for run in sp.runs:
            run.font.size = Pt(FONT_BODY_PT)

    if resume.education:
        _section_header(doc, "Education")
        for edu in resume.education:
            line1 = doc.add_paragraph()
            line1.add_run(edu.institution).bold = True
            if edu.graduation_date:
                line1.add_run(f"    {edu.graduation_date}")
            deg = f"{edu.degree}" + (f" — {edu.field}" if edu.field else "")
            dp = doc.add_paragraph(deg)
            for run in dp.runs:
                run.font.size = Pt(FONT_BODY_PT)

    groups = resume.effective_skill_groups()
    if groups:
        _section_header(doc, "Skills")
        for g in groups:
            p = doc.add_paragraph()
            p.add_run(f"{g.label}: ").bold = True
            p.add_run(", ".join(g.skills))
            for run in p.runs:
                run.font.size = Pt(FONT_BODY_PT)

    if resume.experience:
        _section_header(doc, "Professional Experience")
        for exp in resume.experience:
            h = doc.add_paragraph()
            h.add_run(exp.company).bold = True
            if exp.location:
                h.add_run(f"    {exp.location}")
            sub = doc.add_paragraph(f"{exp.title}    {exp.start_date} – {exp.end_date or 'Present'}")
            for run in sub.runs:
                run.font.size = Pt(FONT_BODY_PT)
            _add_bullets(doc, exp.bullets)

    if resume.projects:
        _section_header(doc, "Projects")
        for proj in resume.projects:
            t = doc.add_paragraph()
            t.add_run(proj.name).bold = True
            if proj.start_date:
                t.add_run(f"    {proj.start_date}" + (f" – {proj.end_date}" if proj.end_date else ""))
            if proj.context:
                cp = doc.add_paragraph(proj.context)
                cp.runs[0].font.size = Pt(FONT_SMALL_PT)
            _add_bullets(doc, proj.bullets)

    if resume.activities:
        _section_header(doc, "Activities & Awards")
        for act in resume.activities:
            doc.add_paragraph(act.title).runs[0].bold = True
            _add_bullets(doc, act.bullets)

    if resume.certifications:
        _section_header(doc, "Certifications")
        for cert in resume.certifications:
            doc.add_paragraph(cert.name)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def build_modern_docx(resume: ResumeData, output_path: Path) -> Path:
    doc = Document()
    _margins(doc)
    accent = RGBColor(*MODERN_ACCENT_RGB)

    name = doc.add_paragraph()
    nr = name.add_run(resume.full_name)
    nr.bold = True
    nr.font.size = Pt(18)
    nr.font.color.rgb = accent

    contact = doc.add_paragraph(_contact_inline(resume))
    for run in contact.runs:
        run.font.size = Pt(FONT_SMALL_PT)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if resume.headline:
        hp = doc.add_paragraph(resume.headline)
        hp.runs[0].bold = True
        hp.runs[0].font.size = Pt(11)

    if resume.summary:
        _section_header(doc, "Profile", color=accent)
        sp = doc.add_paragraph(resume.summary)
        for run in sp.runs:
            run.font.size = Pt(FONT_BODY_PT)

    if resume.experience:
        _section_header(doc, "Experience", color=accent)
        for exp in resume.experience:
            t = doc.add_paragraph()
            tr = t.add_run(exp.title)
            tr.bold = True
            tr.font.color.rgb = accent
            t.add_run(f"    {exp.start_date} – {exp.end_date or 'Present'}")
            loc = doc.add_paragraph(f"{exp.company}" + (f" · {exp.location}" if exp.location else ""))
            for run in loc.runs:
                run.font.size = Pt(FONT_SMALL_PT)
            _add_bullets(doc, exp.bullets)

    if resume.education:
        _section_header(doc, "Education", color=accent)
        for edu in resume.education:
            p = doc.add_paragraph()
            p.add_run(edu.institution).bold = True
            p.add_run(f" — {edu.degree} ({edu.graduation_date})")

    skills_text = ", ".join(resume.skills) if resume.skills else ""
    groups = resume.effective_skill_groups()
    if groups or skills_text:
        _section_header(doc, "Skills", color=accent)
        if groups:
            for g in groups:
                p = doc.add_paragraph()
                p.add_run(f"{g.label}: ").bold = True
                p.add_run(", ".join(g.skills))
        else:
            doc.add_paragraph(skills_text)

    if resume.projects:
        _section_header(doc, "Projects", color=accent)
        for proj in resume.projects:
            doc.add_paragraph(proj.name).runs[0].bold = True
            _add_bullets(doc, proj.bullets)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def build_compact_docx(resume: ResumeData, output_path: Path) -> Path:
    doc = Document()
    _margins(doc, MARGIN_COMPACT_INCH)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nl = header.add_run(resume.full_name)
    nl.bold = True
    nl.font.size = Pt(FONT_NAME_COMPACT_PT)
    header.add_run(f"\t{_contact_inline(resume)}")
    for run in header.runs[1:]:
        run.font.size = Pt(8)

    if resume.summary:
        sp = doc.add_paragraph(resume.summary)
        for run in sp.runs:
            run.font.size = Pt(9)

    if resume.experience:
        _section_header(doc, "Experience")
        for exp in resume.experience:
            line = doc.add_paragraph()
            line.add_run(f"{exp.title}, {exp.company} ").bold = True
            line.add_run(f"({exp.start_date}–{exp.end_date or 'Present'})")
            for run in line.runs:
                run.font.size = Pt(9)
            _add_bullets(doc, exp.bullets, size=9, max_count=4)

    if resume.education:
        _section_header(doc, "Education")
        for edu in resume.education:
            p = doc.add_paragraph(f"{edu.degree}, {edu.institution}")
            for run in p.runs:
                run.font.size = Pt(9)

    if resume.skills:
        _section_header(doc, "Skills")
        p = doc.add_paragraph(" · ".join(resume.skills))
        for run in p.runs:
            run.font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
