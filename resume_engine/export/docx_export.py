from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from resume_engine.schemas.resume import ResumeData
from resume_engine.utils.phone import format_phone_display


def export_docx(resume: ResumeData, output_path: Path) -> Path:
    """Export ATS-friendly single-column DOCX."""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume.full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact_parts = [
        resume.email,
        format_phone_display(resume.phone_country_code, resume.phone),
    ]
    if resume.location:
        contact_parts.append(resume.location)
    if resume.linkedin:
        contact_parts.append(resume.linkedin)
    contact = doc.add_paragraph(" | ".join(contact_parts))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in contact.runs:
        run.font.size = Pt(10)

    # Headline
    if resume.headline:
        headline = doc.add_paragraph(resume.headline)
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in headline.runs:
            run.bold = True
            run.font.size = Pt(11)

    def add_section_header(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # Summary
    if resume.summary:
        add_section_header("Summary")
        summary = doc.add_paragraph(resume.summary)
        for run in summary.runs:
            run.font.size = Pt(10)

    # Experience
    if resume.experience:
        add_section_header("Experience")
        for exp in resume.experience:
            title_line = doc.add_paragraph()
            t_run = title_line.add_run(f"{exp.title} — {exp.company}")
            t_run.bold = True
            t_run.font.size = Pt(10)
            date_line = doc.add_paragraph(f"{exp.start_date} – {exp.end_date}")
            if exp.location:
                date_line.add_run(f" | {exp.location}")
            for run in date_line.runs:
                run.font.size = Pt(9)
            for bullet in exp.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(10)

    # Education
    if resume.education:
        add_section_header("Education")
        for edu in resume.education:
            line = f"{edu.degree}"
            if edu.field:
                line += f" in {edu.field}"
            line += f" — {edu.institution}"
            if edu.graduation_date:
                line += f" ({edu.graduation_date})"
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)

    # Skills
    if resume.skills:
        add_section_header("Skills")
        skills_p = doc.add_paragraph(", ".join(resume.skills))
        for run in skills_p.runs:
            run.font.size = Pt(10)

    # Certifications
    if resume.certifications:
        add_section_header("Certifications")
        for cert in resume.certifications:
            line = cert.name
            if cert.issuer:
                line += f" — {cert.issuer}"
            if cert.date:
                line += f" ({cert.date})"
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
